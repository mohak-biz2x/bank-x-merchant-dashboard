"""Generate Dashboard Insights & Analytics Metrics Summary as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ─── Title Page ───
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Mal Business', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('Dashboard Insights & Analytics\nMetrics Summary & Definitions', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Document Version: ').bold = True
meta.add_run('1.0\n')
meta.add_run('Date: ').bold = True
meta.add_run('June 29, 2026\n')
meta.add_run('Prepared for: ').bold = True
meta.add_run('Mal Bank\n')
meta.add_run('Prototype: ').bold = True
meta.add_run('https://mohak-biz2x.github.io/bank-x-merchant-dashboard/dashboard-prototypes-index.html')

doc.add_page_break()

# ─── Section 1: Overview ───
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'The Dashboard Insights & Analytics module adds visual data representations to the Merchant Dashboard, '
    'providing merchants with at-a-glance visibility into their financing activity, portfolio health, and '
    'counterparty risk. The module adapts to the merchant\'s financing product (Payable or Receivable) and activity level.'
)

doc.add_heading('Prototype Variants', level=2)
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Variant'
hdr[1].text = 'Description'
hdr[2].text = 'Key Data'
data = [
    ('Payable Financing (Full Data)', 'Active merchant using payable invoice financing', '6 suppliers, 82 invoices, 64% utilised'),
    ('Receivable Financing (Full Data)', 'Active merchant using receivable invoice financing', '5 buyers, 50 invoices, 64% utilised'),
    ('Low Activity', 'Merchant who just started (1-2 months of data)', '1 supplier, 5 invoices, 6% utilised'),
    ('No Activity (Empty State)', 'Merchant with limit just enabled, zero transactions', '0 invoices, 0% utilised'),
]
for i, row_data in enumerate(data):
    row = table.rows[i + 1].cells
    row[0].text = row_data[0]
    row[1].text = row_data[1]
    row[2].text = row_data[2]

doc.add_paragraph()

# ─── Section 2: Metrics & Definitions ───
doc.add_heading('2. Metrics & Definitions', level=1)

metrics = [
    {
        'title': '2.1 Credit Utilisation',
        'chart_type': 'Area/Line chart with reference line',
        'definition': 'Shows how much of the approved credit limit has been drawn down over time.',
        'formula': 'Utilised Amount ÷ Approved Limit × 100',
        'time_periods': '3 months, 6 months, 12 months',
        'visibility': 'Payable: amber payable line only. Receivable: green receivable line only.',
        'business_value': 'Helps merchant plan financing capacity and avoid hitting limit.',
        'notes': 'A dashed red reference line marks the Approved Limit (e.g., AED 8,000,000).'
    },
    {
        'title': '2.2 Invoice Count Trend',
        'chart_type': 'Line chart with area fill',
        'definition': 'Count of invoices submitted per month over the selected time period.',
        'formula': 'Count of invoices submitted in each calendar month',
        'time_periods': '3 months, 6 months, 12 months',
        'visibility': 'Both Payable and Receivable roles.',
        'business_value': 'Identifies growth trends, seasonality, and submission patterns.',
        'notes': ''
    },
    {
        'title': '2.3 Payment Terms Distribution',
        'chart_type': 'Donut chart',
        'definition': 'Breakdown of all invoices by their payment tenor (days until due).',
        'formula': 'Count of invoices per tenor bucket (30, 60, 90, 120 days)',
        'time_periods': 'All-time (no time selector)',
        'visibility': 'Both Payable and Receivable roles.',
        'business_value': 'Shows the tenor mix — helps assess cash flow timing and financing cost exposure.',
        'notes': 'Center of donut displays total invoice count.'
    },
    {
        'title': '2.4 Supplier / Buyer Distribution',
        'chart_type': 'Donut chart',
        'definition': 'Distribution of total financed amount across counterparties.',
        'formula': 'Invoice amount per counterparty ÷ Total financed amount × 100',
        'time_periods': 'All-time (no time selector)',
        'visibility': 'Payable: "Supplier Distribution". Receivable: "Buyer Distribution".',
        'business_value': 'Visualizes counterparty concentration — identifies over-reliance on single entities.',
        'notes': 'Top 5 counterparties shown individually; remainder grouped as "Others". Center shows total AED amount.'
    },
    {
        'title': '2.5 Approval Ratio',
        'chart_type': 'Score metric with progress bar',
        'definition': 'Percentage of invoices approved out of all decisioned (approved + rejected) invoices.',
        'formula': 'Approved Invoices ÷ (Approved + Rejected) × 100',
        'time_periods': 'All-time cumulative (not month-over-month)',
        'visibility': 'Both Payable and Receivable roles.',
        'business_value': 'Single health metric showing invoice quality — high ratio means fewer rejections.',
        'notes': 'Excludes Pending and Refer invoices (not yet decisioned). Thresholds: ≥90% = Healthy (green), 70-89% = Average (amber), <70% = Low (red). Designed as cumulative to handle irregular submission patterns.'
    },
    {
        'title': '2.6 Concentration Risk Score',
        'chart_type': 'Score metric with gradient bar',
        'definition': 'Measures how dependent the portfolio is on a small number of counterparties.',
        'formula': 'Derived from HHI (Herfindahl-Hirschman Index) — sum of squared market shares',
        'time_periods': 'All-time (no time selector)',
        'visibility': 'Payable: supplier concentration. Receivable: buyer concentration.',
        'business_value': 'Alerts merchant to concentration risk — over-reliance on one counterparty increases default exposure.',
        'notes': 'Score 0-100. Thresholds: 0-40 = Low (green), 41-65 = Moderate (amber), 66-100 = High (red). Supporting metrics: Top 1 %, Top 3 %, HHI Index. HHI range: 0-10,000 (above 2,500 = high concentration).'
    },
]

for m in metrics:
    doc.add_heading(m['title'], level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Light Grid Accent 1'
    fields = [
        ('Chart Type', m['chart_type']),
        ('Definition', m['definition']),
        ('Formula', m['formula']),
        ('Time Periods', m['time_periods']),
        ('Role Visibility', m['visibility']),
        ('Business Value', m['business_value']),
        ('Notes', m['notes']),
    ]
    for i, (label, value) in enumerate(fields):
        row = table.rows[i].cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].bold = True if row[0].paragraphs[0].runs else None
        row[1].text = value
    doc.add_paragraph()

# ─── Section 3: Role-Based Visibility ───
doc.add_heading('3. Role-Based Visibility Matrix', level=1)
table = doc.add_table(rows=8, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['Metric', 'Payable', 'Receivable', 'Both']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

visibility_data = [
    ('Credit Utilisation', '✓ (payable only)', '✓ (receivable only)', '✓ (both series)'),
    ('Invoice Count Trend', '✓', '✓', '✓'),
    ('Payment Terms Distribution', '✓', '✓', '✓'),
    ('Supplier Distribution', '✓', '—', '✓'),
    ('Buyer Distribution', '—', '✓', '—'),
    ('Approval Ratio', '✓', '✓', '✓'),
    ('Concentration Risk', '✓ (supplier)', '✓ (buyer)', '✓ (supplier)'),
]
for i, row_data in enumerate(visibility_data):
    for j, val in enumerate(row_data):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph('Excluded roles: supplier-only, premium-buyer, premium-buyer-supplier do not see the analytics section.').italic = True

# ─── Section 4: Empty & Low-Activity States ───
doc.add_heading('4. Empty & Low-Activity States', level=1)
table = doc.add_table(rows=3, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'State'
table.rows[0].cells[1].text = 'Behavior'
table.rows[1].cells[0].text = 'No Activity'
table.rows[1].cells[1].text = 'Charts replaced with empty state card: "No activity yet — Your insights will appear here once you start submitting invoices and adding suppliers."'
table.rows[2].cells[0].text = 'Low Activity (< 3 months)'
table.rows[2].cells[1].text = 'Charts render with available data points; 6M and 12M time selectors are disabled; info banner shown: "Your analytics will become more detailed as your financing activity grows."'

# ─── Section 5: Interactive Features ───
doc.add_heading('5. Interactive Features', level=1)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'
table.rows[0].cells[0].text = 'Feature'
table.rows[0].cells[1].text = 'Behavior'
features = [
    ('Time Period Selector', 'Segmented buttons (3M / 6M / 12M) that re-render charts with selected data range'),
    ('Tooltips', 'Hover/tap on any data point shows detailed values (month, amount, percentage)'),
    ('Responsive Layout', '2-column grid on desktop, single column on mobile (<1024px)'),
    ('Card Hover Effect', 'Subtle indigo border glow on hover'),
]
for i, (feat, desc) in enumerate(features):
    table.rows[i + 1].cells[0].text = feat
    table.rows[i + 1].cells[1].text = desc

# ─── Section 6: Data Source Notes ───
doc.add_heading('6. Data Source Notes', level=1)
doc.add_paragraph(
    'All data in the prototype is mock/hardcoded for demonstration purposes. In production:'
)
sources = [
    'Credit utilisation data — Loan Management System (LMS)',
    'Invoice counts and statuses — Invoice Processing Engine',
    'Supplier/Buyer data — KYB/Onboarding module',
    'Pricing/cost data — Pricing Rule Engine',
]
for s in sources:
    doc.add_paragraph(s, style='List Bullet')

# ─── Save ───
output_path = r'c:\Merchant Dashbaord\docs\Dashboard_Insights_Analytics_Metrics_Summary.docx'
doc.save(output_path)
print(f'Word document saved to: {output_path}')
